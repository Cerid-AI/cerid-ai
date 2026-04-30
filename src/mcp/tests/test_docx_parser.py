# Copyright (c) 2026 Cerid AI. All rights reserved.
# SPDX-License-Identifier: Apache-2.0

"""Tests for the DOCX parser (Workstream E Phase 2b.7)."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from core.ingest.chunkers import chunk_elements
from core.ingest.parsers.docx_parser import parse_docx


def _build_docx(path: Path, builder) -> Path:
    """Build a docx via a callback that takes the python-docx Document."""
    doc = Document()
    builder(doc)
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------


def test_parse_docx_emits_section_per_heading_group(tmp_path):
    def build(d):
        d.add_heading("Top", level=1)
        d.add_paragraph("Intro paragraph.")
        d.add_heading("Sub A", level=2)
        d.add_paragraph("Body of A.")
        d.add_heading("Sub B", level=2)
        d.add_paragraph("Body of B.")

    p = _build_docx(tmp_path / "doc.docx", build)
    elements = parse_docx(p)
    sections = [el for el in elements if el["element_type"] == "MarkdownSection"]
    assert len(sections) == 3


def test_parse_docx_heading_path_tracks_hierarchy(tmp_path):
    def build(d):
        d.add_heading("Top", level=1)
        d.add_heading("Sub", level=2)
        d.add_heading("Leaf", level=3)
        d.add_paragraph("Leaf body.")

    p = _build_docx(tmp_path / "doc.docx", build)
    elements = parse_docx(p)
    leaf = next(
        el for el in elements
        if el["element_type"] == "MarkdownSection"
        and "Leaf body" in el["text"]
    )
    assert leaf["metadata"]["heading_path"] == ["Top", "Sub", "Leaf"]
    assert leaf["metadata"]["level"] == 3


def test_parse_docx_heading_path_resets_when_jumping_back_up(tmp_path):
    """A new H1 should clear the H2/H3 stack."""
    def build(d):
        d.add_heading("Section A", level=1)
        d.add_heading("Sub A1", level=2)
        d.add_paragraph("A1 body")
        d.add_heading("Section B", level=1)  # back up to H1
        d.add_paragraph("B body")

    p = _build_docx(tmp_path / "doc.docx", build)
    elements = parse_docx(p)
    b_section = next(
        el for el in elements
        if el["element_type"] == "MarkdownSection"
        and "B body" in el["text"]
    )
    # Section B's path is just ['Section B'], NOT ['Section A', 'Section B']
    # and NOT carrying the prior 'Sub A1'.
    assert b_section["metadata"]["heading_path"] == ["Section B"]


def test_parse_docx_extracts_tables_separately(tmp_path):
    def build(d):
        d.add_heading("Data", level=1)
        d.add_paragraph("Here's some data.")
        t = d.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "name"
        t.rows[0].cells[1].text = "value"
        t.rows[1].cells[0].text = "alice"
        t.rows[1].cells[1].text = "42"

    p = _build_docx(tmp_path / "doc.docx", build)
    elements = parse_docx(p)
    tables = [el for el in elements if el["element_type"] == "Table"]
    assert len(tables) == 1
    tbl = tables[0]
    assert "| name | value |" in tbl["text"]
    assert "| alice | 42 |" in tbl["text"]
    # Heading_path on the table — same Data section
    assert tbl["metadata"]["heading_path"] == ["Data"]
    assert tbl["metadata"]["rows"] == [["name", "value"], ["alice", "42"]]


def test_parse_docx_preserves_document_order_paragraphs_and_tables(tmp_path):
    """Sections and tables should appear in the order they're written."""
    def build(d):
        d.add_heading("Section A", level=1)
        d.add_paragraph("Para A1")
        t1 = d.add_table(rows=1, cols=1)
        t1.rows[0].cells[0].text = "Table 1 cell"
        d.add_heading("Section B", level=1)
        d.add_paragraph("Para B1")
        t2 = d.add_table(rows=1, cols=1)
        t2.rows[0].cells[0].text = "Table 2 cell"

    p = _build_docx(tmp_path / "doc.docx", build)
    elements = parse_docx(p)
    sequence = []
    for el in elements:
        if el["element_type"] == "MarkdownSection":
            sequence.append(("section", el["metadata"]["heading_path"][-1]))
        elif el["element_type"] == "Table":
            sequence.append(("table", el["metadata"]["rows"][0][0]))
    # Document-order: A section → A table → B section → B table
    assert sequence == [
        ("section", "Section A"),
        ("table", "Table 1 cell"),
        ("section", "Section B"),
        ("table", "Table 2 cell"),
    ]


def test_parse_docx_handles_title_style_as_h1(tmp_path):
    def build(d):
        # add_heading(level=0) writes a 'Title' style paragraph
        d.add_heading("Doc Title", level=0)
        d.add_paragraph("Body under the title.")

    p = _build_docx(tmp_path / "doc.docx", build)
    elements = parse_docx(p)
    sections = [el for el in elements if el["element_type"] == "MarkdownSection"]
    assert sections
    assert sections[0]["metadata"]["heading_path"] == ["Doc Title"]


def test_parse_docx_skips_empty_paragraphs(tmp_path):
    def build(d):
        d.add_heading("Top", level=1)
        d.add_paragraph("")  # empty
        d.add_paragraph("Real body.")
        d.add_paragraph("   ")  # whitespace only

    p = _build_docx(tmp_path / "doc.docx", build)
    elements = parse_docx(p)
    section = next(el for el in elements if el["element_type"] == "MarkdownSection")
    # Only "Real body." survives
    assert "Real body" in section["text"]


def test_parse_docx_paragraph_before_first_heading_emits_with_empty_path(tmp_path):
    """Real-world docs sometimes start with a paragraph before any heading."""
    def build(d):
        d.add_paragraph("Preamble text before any heading.")
        d.add_heading("Top", level=1)
        d.add_paragraph("Under the heading.")

    p = _build_docx(tmp_path / "doc.docx", build)
    elements = parse_docx(p)
    preamble = next(
        el for el in elements
        if el["element_type"] == "MarkdownSection"
        and "Preamble" in el["text"]
    )
    assert preamble["metadata"]["heading_path"] == []
    assert preamble["metadata"]["level"] == 0


def test_parse_docx_raises_on_missing_file(tmp_path):
    with pytest.raises(FileNotFoundError):
        parse_docx(tmp_path / "missing.docx")


def test_parse_docx_swallows_corrupt_file(tmp_path):
    p = tmp_path / "corrupt.docx"
    p.write_bytes(b"PK-broken garbage that python-docx will reject")
    assert parse_docx(p) == []


def test_parse_docx_empty_document_returns_empty(tmp_path):
    p = _build_docx(tmp_path / "empty.docx", lambda d: None)
    assert parse_docx(p) == []


# ---------------------------------------------------------------------------
# End-to-end dispatch — DOCX reuses Markdown + PDF Table strategies
# ---------------------------------------------------------------------------


def test_chunk_elements_dispatches_docx_via_markdown_strategy(tmp_path):
    def build(d):
        d.add_heading("Top", level=1)
        d.add_heading("Sub", level=2)
        d.add_paragraph("Leaf body.")

    p = _build_docx(tmp_path / "doc.docx", build)
    elements = parse_docx(p)
    chunks = chunk_elements(elements)
    section_chunks = [
        c for c in chunks if c["metadata"]["element_type"] == "MarkdownSection"
    ]
    assert section_chunks
    # Markdown strategy prepends the heading breadcrumb
    assert any(c["text"].startswith("Top > Sub") for c in section_chunks)


def test_chunk_elements_dispatches_docx_table_via_pdf_strategy(tmp_path):
    def build(d):
        d.add_heading("Data", level=1)
        t = d.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "x"
        t.rows[0].cells[1].text = "y"

    p = _build_docx(tmp_path / "doc.docx", build)
    elements = parse_docx(p)
    chunks = chunk_elements(elements)
    table_chunks = [c for c in chunks if c["metadata"]["element_type"] == "Table"]
    assert table_chunks
    md = table_chunks[0]["metadata"]
    # heading_path metadata propagates from parser through the Table strategy
    assert md["heading_path"] == ["Data"]
