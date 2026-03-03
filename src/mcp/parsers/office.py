# Copyright (c) 2026 Justin Michaels. All rights reserved.
# SPDX-License-Identifier: Apache-2.0

"""Office document parsers — DOCX and XLSX."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from parsers.registry import _MAX_TEXT_CHARS, logger, register_parser


@register_parser([".docx"])
def parse_docx(file_path: str) -> dict[str, Any]:
    import docx

    try:
        doc = docx.Document(file_path)
    except Exception as e:
        raise ValueError(
            f"Failed to read DOCX '{Path(file_path).name}': {e}. "
            f"File may be corrupted or not a valid .docx file."
        ) from e

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    table_texts = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append("\t".join(cells))
        if rows:
            table_texts.append("\n".join(rows))

    parts = paragraphs
    if table_texts:
        parts.append("\n--- Tables ---")
        parts.extend(table_texts)

    text = "\n\n".join(parts)
    return {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": "docx",
        "page_count": None,
    }


@register_parser([".xlsx"])
def parse_xlsx(file_path: str) -> dict[str, Any]:
    """Parse XLSX with header auto-detection and Markdown table formatting."""
    from openpyxl import load_workbook

    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)
    except Exception as e:
        raise ValueError(
            f"Failed to read XLSX '{Path(file_path).name}': {e}. "
            f"File may be corrupted or not a valid .xlsx file."
        ) from e

    sheet_names = list(wb.sheetnames)
    sheets_text = []
    all_columns: list[str] = []
    total_rows = 0
    truncated = False

    for sheet_name in sheet_names:
        ws = wb[sheet_name]
        raw_rows = []
        for row in ws.iter_rows(values_only=True):
            cell_values = [str(c) if c is not None else "" for c in row]
            if any(cell_values):
                raw_rows.append(cell_values)

        if not raw_rows:
            continue

        total_rows += len(raw_rows)

        header_idx = 0
        for idx, row in enumerate(raw_rows[:5]):
            non_empty = sum(1 for c in row if c.strip())
            if len(row) > 0 and non_empty / len(row) > 0.5:
                header_idx = idx
                break

        header = raw_rows[header_idx]
        all_columns.extend([c.strip() for c in header if c.strip()])
        data_rows = raw_rows[header_idx + 1:]

        if len(data_rows) > 5000:
            logger.warning(
                f"XLSX '{Path(file_path).name}' sheet '{sheet_name}' has "
                f"{len(data_rows)} rows, truncating to 5000"
            )
            data_rows = data_rows[:5000]
            truncated = True

        md_lines = [f"--- Sheet: {sheet_name} ---"]
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in data_rows:
            # Pad or trim row to match header column count
            padded = row + [""] * max(0, len(header) - len(row))
            md_lines.append("| " + " | ".join(padded[:len(header)]) + " |")
        sheets_text.append("\n".join(md_lines))

    wb.close()

    text = "\n\n".join(sheets_text)

    result: dict[str, Any] = {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": "xlsx",
        "page_count": len(sheet_names),
        "row_count": total_rows,
    }
    if all_columns:
        seen: set = set()
        unique_cols = []
        for c in all_columns:
            if c not in seen:
                seen.add(c)
                unique_cols.append(c)
        result["columns"] = json.dumps(unique_cols[:50])
    if truncated:
        result["truncated"] = True

    return result
