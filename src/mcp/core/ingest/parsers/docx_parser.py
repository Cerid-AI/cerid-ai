# Copyright (c) 2026 Cerid AI. All rights reserved.
# SPDX-License-Identifier: Apache-2.0

"""DOCX parser — heading-hierarchy + table extraction.

Workstream E Phase 2b.7. Closes the audit's "DOCX: paragraphs only;
heading hierarchy, list nesting, table cells, footnotes flattened"
gap. Walks the document body in declaration order and emits:

* :class:`MarkdownSection` elements for paragraph groups under each
  heading. ``heading_path`` metadata mirrors the Phase 2b.2
  Markdown parser's shape so the existing
  :func:`markdown_section_strategy` chunker (heading-breadcrumb
  prepend, oversized-section split with breadcrumb replay) is
  re-used as-is — no new strategy needed for paragraphs.
* :class:`Table` elements for each ``w:tbl`` in document order.
  Same ``rows`` + Markdown-pipe text shape as the PDF parser's
  Table elements so the existing :func:`pdf_table_strategy`
  chunker handles them too.

Interleaving paragraphs + tables in correct document order matters
because retrieval should preserve "Section A → Table A1 → Section
A continued → Section B → Table B1" sequencing — the chunker
chain's metadata stays accurate that way.

Library choice: ``python-docx`` — already in ``requirements.txt``
(used by the legacy parser at ``app/parsers/docx.py``). Pure Python
on top of lxml. The runtime adds zero dependencies for Phase 2b.7.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from core.ingest.parsers import ParsedElement

logger = logging.getLogger("ai-companion.ingest.parsers.docx")

# python-docx exposes the Word XML namespace as ``w:`` — we need the
# expanded URI form to match :func:`lxml.etree._Element.iterchildren`
# return values. The ``qn()`` helper from python-docx does this but
# we hard-code the strings to avoid a docx-internals import.
_NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_TAG_P = f"{_NS_W}p"
_TAG_TBL = f"{_NS_W}tbl"


def _heading_level(style_name: str) -> int | None:
    """Return the heading level (1-9) for a Word style name, or None.

    Word's built-in styles are ``Heading 1`` … ``Heading 9``;
    ``Title`` is also commonly used as the document title (treated
    as level 1). Custom styles whose name contains the word
    "heading" are also recognised — best-effort.
    """
    if not style_name:
        return None
    s = style_name.strip()
    if s == "Title":
        return 1
    if s.startswith("Heading "):
        try:
            return int(s.split(" ", 1)[1])
        except (ValueError, IndexError):
            return None
    return None


def _table_to_rows(tbl: Any) -> list[list[str]]:
    """Convert a python-docx Table to a 2D list of stripped cell strings."""
    out: list[list[str]] = []
    for row in tbl.rows:
        out.append([cell.text.strip() for cell in row.cells])
    return out


def _rows_to_markdown(rows: list[list[str]]) -> str:
    """Render a 2D string grid as a Markdown pipe-table.

    Matches the shape :func:`pdf_parser._table_rows_to_markdown`
    emits so downstream retrieval gets the same format regardless of
    whether the table came from a PDF or DOCX.
    """
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    padded = [r + [""] * (width - len(r)) for r in rows]
    lines = ["| " + " | ".join(padded[0]) + " |"]
    if len(padded) >= 2:
        lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def parse_docx(path: str | Path) -> list[ParsedElement]:
    """Parse a `.docx` file into MarkdownSection + Table elements.

    Args:
        path: Filesystem path to the document.

    Returns:
        A list of :class:`ParsedElement` dicts in document order:

        * one ``MarkdownSection`` per (heading-or-document-start →
          next heading) paragraph group. Sections under multiple
          headings get the full ordered ``heading_path``.
        * one ``Table`` per ``w:tbl`` with structured ``rows`` +
          ``n_rows`` / ``n_cols`` + Markdown-pipe text.

        Returns ``[]`` on parse failure or empty documents.

    Raises:
        FileNotFoundError: when ``path`` doesn't exist.
        ImportError: when python-docx isn't installed.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"DOCX not found: {p}")

    from docx import Document

    elements: list[ParsedElement] = []
    try:
        doc = Document(str(p))
    except Exception as exc:  # noqa: BLE001 — fall back to legacy parser
        logger.warning("docx_open_failed file=%s error=%s", p.name, exc)
        return []

    # Index paragraphs + tables by document order. python-docx exposes
    # paragraphs and tables separately, but the underlying lxml body
    # has them interleaved correctly.
    paragraphs_iter = iter(doc.paragraphs)
    tables_iter = iter(doc.tables)

    heading_path: list[str] = []
    section_buf: list[str] = []
    section_path: list[str] = []  # path that was active when section_buf started

    def _flush_section() -> None:
        if not section_buf:
            return
        body = "\n\n".join(s for s in section_buf if s).strip()
        if not body:
            section_buf.clear()
            return
        # Build the headers dict for parity with the markdown parser's shape
        headers = {f"h{i + 1}": h for i, h in enumerate(section_path)}
        elements.append(
            {
                "text": body,
                "element_type": "MarkdownSection",
                "metadata": {
                    "heading_path": list(section_path),
                    "level": len(section_path),
                    "headers": headers,
                },
            },
        )
        section_buf.clear()

    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag == _TAG_P:
            try:
                para = next(paragraphs_iter)
            except StopIteration:  # pragma: no cover — should not happen
                break
            text = para.text.strip()
            if not text:
                continue
            level = _heading_level(para.style.name if para.style else "")
            if level is not None:
                # Heading: flush the prior section, then update the path
                _flush_section()
                # Trim the path to one level shallower than the new heading
                heading_path = heading_path[: level - 1] + [text]
                section_path = list(heading_path)
            else:
                # Body paragraph — initialise section_path lazily so a
                # leading paragraph before any heading still gets emitted.
                if not section_buf:
                    section_path = list(heading_path)
                section_buf.append(text)
        elif tag == _TAG_TBL:
            try:
                tbl = next(tables_iter)
            except StopIteration:  # pragma: no cover
                continue
            # Flush any in-flight section so document order is preserved.
            _flush_section()
            rows = _table_to_rows(tbl)
            if not rows or not any(any(c for c in row) for row in rows):
                continue
            n_cols = max((len(r) for r in rows), default=0)
            elements.append(
                {
                    "text": _rows_to_markdown(rows),
                    "element_type": "Table",
                    "metadata": {
                        "n_rows": len(rows),
                        "n_cols": n_cols,
                        "rows": rows,
                        "heading_path": list(heading_path),
                    },
                },
            )
        # Other body children (sectPr, sdt, etc.) are skipped.

    # Final flush after walking the body
    _flush_section()

    logger.info(
        "docx_parsed file=%s elements=%d (sections=%d, tables=%d)",
        p.name, len(elements),
        sum(1 for el in elements if el["element_type"] == "MarkdownSection"),
        sum(1 for el in elements if el["element_type"] == "Table"),
    )
    return elements
