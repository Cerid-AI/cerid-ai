"""
Extensible file parser registry.

Parsers are registered by file extension via the @register_parser decorator.
To add a new parser (e.g. Docling): define a function and decorate it.
To override an existing parser: register the same extension — last wins.

Phase 8B additions:
- .eml / .mbox — Email parsing (headers + body + attachment list)
- .epub — E-book text extraction
- .rtf — Rich text format parsing
- Enhanced XLSX/CSV/PDF parsers (schema summary, delimiter detection, form fields)
"""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any, Callable, Dict, List

logger = logging.getLogger("ai-companion.parsers")

# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------
PARSER_REGISTRY: Dict[str, Callable[[str], Dict[str, Any]]] = {}


def register_parser(extensions: List[str]):
    """Decorator that maps file extensions to a parser function."""
    def decorator(func: Callable[[str], Dict[str, Any]]):
        for ext in extensions:
            PARSER_REGISTRY[ext.lower()] = func
        return func
    return decorator


def parse_file(file_path: str) -> Dict[str, Any]:
    """
    Parse a file and return its text content + metadata.

    Returns:
        {"text": str, "file_type": str, "page_count": int | None}

    Raises:
        ValueError: unsupported extension or parse failure with clear message
        FileNotFoundError: file does not exist
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if path.stat().st_size == 0:
        raise ValueError(f"File is empty (0 bytes): {path.name}")

    ext = path.suffix.lower()
    parser = PARSER_REGISTRY.get(ext)
    if not parser:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: {sorted(PARSER_REGISTRY.keys())}"
        )
    return parser(file_path)


# ---------------------------------------------------------------------------
# Built-in parsers
# ---------------------------------------------------------------------------

# Maximum text output size to prevent memory issues with huge files
_MAX_TEXT_CHARS = 2_000_000  # ~2MB of text


@register_parser([".pdf"])
def parse_pdf(file_path: str) -> Dict[str, Any]:
    """
    Parse PDFs using pdfplumber for structure-aware extraction.

    Preserves table layouts as Markdown-style tables, maintaining the spatial
    relationship between labels and values (critical for tax forms, financial
    statements, and any document with structured grids).

    Falls back to raw text extraction for pages without tables.
    """
    import pdfplumber

    try:
        pdf = pdfplumber.open(file_path)
    except Exception as e:
        raise ValueError(
            f"Failed to read PDF '{Path(file_path).name}': {e}. "
            f"File may be corrupted, password-protected, or not a valid PDF."
        ) from e

    page_count = len(pdf.pages)
    pages = []
    table_count = 0

    try:
        for i, page in enumerate(pdf.pages):
            try:
                page_parts = []

                # Extract tables first — they contain structured data that
                # plain text extraction would jumble
                tables = page.find_tables()
                if tables:
                    # Get bounding boxes of all tables so we can exclude them
                    # from plain text extraction (avoids duplicate content)
                    table_bboxes = [t.bbox for t in tables]

                    for table in tables:
                        table_count += 1
                        rows = table.extract()
                        if rows:
                            # Format as Markdown table for structure preservation
                            md_rows = []
                            for row in rows:
                                cells = [
                                    (cell or "").strip().replace("\n", " ")
                                    for cell in row
                                ]
                                md_rows.append("| " + " | ".join(cells) + " |")
                            # Add header separator after first row
                            if len(md_rows) > 1:
                                col_count = len(rows[0]) if rows[0] else 1
                                md_rows.insert(1, "| " + " | ".join(["---"] * col_count) + " |")
                            page_parts.append("\n".join(md_rows))

                    # Extract non-table text by cropping around table regions
                    # This prevents duplicating content that's already in tables
                    try:
                        filtered = page
                        for bbox in table_bboxes:
                            # Crop out each table region from the page
                            # pdfplumber bbox: (x0, top, x1, bottom)
                            filtered = filtered.outside_bounding_box(bbox)
                        plain_text = filtered.extract_text()
                    except Exception:
                        # If cropping fails, fall back to full page text
                        # (may duplicate some table content — acceptable)
                        plain_text = page.extract_text()

                    if plain_text and plain_text.strip():
                        page_parts.insert(0, plain_text.strip())
                else:
                    # No tables on this page — standard text extraction
                    plain_text = page.extract_text()
                    if plain_text and plain_text.strip():
                        page_parts.append(plain_text.strip())

                if page_parts:
                    pages.append("\n\n".join(page_parts))
            except Exception as e:
                logger.warning(f"PDF page {i+1}/{page_count} failed to extract: {e}")
    finally:
        pdf.close()
    text = "\n\n".join(pages)

    # Phase 8B: Extract form fields (AcroForm) if present
    form_fields = []
    try:
        pdf2 = pdfplumber.open(file_path)
        try:
            for page in pdf2.pages:
                annots = page.annots or []
                for annot in annots:
                    if isinstance(annot, dict):
                        field_name = annot.get("T", "")
                        field_value = annot.get("V", "")
                        if field_name or field_value:
                            form_fields.append(
                                f"{field_name}: {field_value}" if field_name
                                else str(field_value)
                            )
        finally:
            pdf2.close()
    except Exception:
        pass  # form field extraction is best-effort

    if form_fields:
        text += "\n\n--- Form Fields ---\n" + "\n".join(form_fields[:100])

    # Detect image-only PDFs (pages exist but no text extracted)
    if not text.strip() and page_count > 0:
        raise ValueError(
            f"No text extracted from PDF '{Path(file_path).name}' "
            f"({page_count} pages). This is likely a scanned/image-only PDF. "
            f"OCR support (e.g. Docling) is needed to process this file."
        )

    if table_count:
        logger.info(
            f"PDF '{Path(file_path).name}': {page_count} pages, "
            f"{table_count} tables extracted as Markdown"
        )

    result: Dict[str, Any] = {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": "pdf",
        "page_count": page_count,
        "table_count": table_count,
    }
    if form_fields:
        result["form_field_count"] = len(form_fields)

    return result


@register_parser([".docx"])
def parse_docx(file_path: str) -> Dict[str, Any]:
    import docx

    try:
        doc = docx.Document(file_path)
    except Exception as e:
        raise ValueError(
            f"Failed to read DOCX '{Path(file_path).name}': {e}. "
            f"File may be corrupted or not a valid .docx file."
        ) from e

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Extract tables (often contain critical data missed by paragraph-only extraction)
    table_texts = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append("\t".join(cells))
        if rows:
            table_texts.append("\n".join(rows))

    parts = paragraphs
    if table_texts:
        parts.append("\n--- Tables ---")
        parts.extend(table_texts)

    text = "\n\n".join(parts)
    return {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": "docx",
        "page_count": None,
    }


@register_parser([".xlsx"])
def parse_xlsx(file_path: str) -> Dict[str, Any]:
    """
    Parse XLSX with header detection and schema summary.

    Phase 8B enhancement:
    - Auto-detects header row (first row with text in >50% of columns)
    - Extracts column names as metadata keywords
    - Truncation warning if sheet has >5000 rows
    - Formats as Markdown table for structure preservation
    """
    from openpyxl import load_workbook

    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)
    except Exception as e:
        raise ValueError(
            f"Failed to read XLSX '{Path(file_path).name}': {e}. "
            f"File may be corrupted or not a valid .xlsx file."
        ) from e

    sheet_names = list(wb.sheetnames)
    sheets_text = []
    all_columns: List[str] = []
    total_rows = 0
    truncated = False

    for sheet_name in sheet_names:
        ws = wb[sheet_name]
        raw_rows = []
        for row in ws.iter_rows(values_only=True):
            cell_values = [str(c) if c is not None else "" for c in row]
            if any(cell_values):
                raw_rows.append(cell_values)

        if not raw_rows:
            continue

        total_rows += len(raw_rows)

        # Detect header row: first row with text in >50% of columns
        header_idx = 0
        for idx, row in enumerate(raw_rows[:5]):
            non_empty = sum(1 for c in row if c.strip())
            if len(row) > 0 and non_empty / len(row) > 0.5:
                header_idx = idx
                break

        header = raw_rows[header_idx]
        all_columns.extend([c.strip() for c in header if c.strip()])
        data_rows = raw_rows[header_idx + 1:]

        # Truncate large sheets
        if len(data_rows) > 5000:
            logger.warning(
                f"XLSX '{Path(file_path).name}' sheet '{sheet_name}' has "
                f"{len(data_rows)} rows, truncating to 5000"
            )
            data_rows = data_rows[:5000]
            truncated = True

        # Format as Markdown table
        md_lines = [f"--- Sheet: {sheet_name} ---"]
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in data_rows:
            # Pad or trim row to match header column count
            padded = row + [""] * max(0, len(header) - len(row))
            md_lines.append("| " + " | ".join(padded[:len(header)]) + " |")
        sheets_text.append("\n".join(md_lines))

    wb.close()

    text = "\n\n".join(sheets_text)

    result: Dict[str, Any] = {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": "xlsx",
        "page_count": len(sheet_names),
        "row_count": total_rows,
    }
    if all_columns:
        # Deduplicate while preserving order
        seen: set = set()
        unique_cols = []
        for c in all_columns:
            if c not in seen:
                seen.add(c)
                unique_cols.append(c)
        result["columns"] = json.dumps(unique_cols[:50])
    if truncated:
        result["truncated"] = True

    return result


@register_parser([".csv", ".tsv"])
def parse_csv(file_path: str) -> Dict[str, Any]:
    """
    Parse CSV/TSV with auto-delimiter detection and schema summary.

    Phase 8B enhancements:
    - Auto-detect delimiter (csv.Sniffer) — handles comma, tab, semicolon, pipe
    - .tsv extension support
    - Schema summary: column names, inferred types, row count
    - Sample rows (first 5) preserved in metadata for context
    """
    import csv as csv_module

    import pandas as pd

    fname = Path(file_path).name
    ext = Path(file_path).suffix.lower()

    # Auto-detect delimiter
    delimiter = "\t" if ext == ".tsv" else ","
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            sample = f.read(8192)
        try:
            dialect = csv_module.Sniffer().sniff(sample, delimiters=",;\t|")
            delimiter = dialect.delimiter
        except csv_module.Error:
            pass  # keep default
    except Exception:
        pass

    try:
        try:
            df = pd.read_csv(file_path, encoding="utf-8", sep=delimiter)
        except UnicodeDecodeError:
            df = pd.read_csv(file_path, encoding="latin-1", sep=delimiter)
    except Exception as e:
        raise ValueError(
            f"Failed to read CSV '{fname}': {e}. "
            f"File may be corrupted or not a valid CSV."
        ) from e

    row_count = len(df)
    columns = list(df.columns)
    truncated = False

    if row_count > 5000:
        logger.warning(
            f"CSV '{fname}' has {row_count} rows, truncating to first 5000 for ingestion"
        )
        df = df.head(5000)
        truncated = True

    # Build schema summary for metadata
    type_map = {}
    for col in columns:
        dtype = str(df[col].dtype)
        if "int" in dtype:
            type_map[col] = "integer"
        elif "float" in dtype:
            type_map[col] = "number"
        elif "datetime" in dtype:
            type_map[col] = "datetime"
        elif "bool" in dtype:
            type_map[col] = "boolean"
        else:
            type_map[col] = "text"

    # Build structured text with schema header and sample rows
    schema_lines = [f"Schema: {len(columns)} columns, {row_count} rows"]
    schema_lines.append("Columns: " + ", ".join(f"{c} ({type_map.get(c, 'text')})" for c in columns[:30]))
    if len(columns) > 30:
        schema_lines.append(f"  ... and {len(columns) - 30} more columns")

    # Sample rows (first 5) for context
    sample_df = df.head(5)
    sample_text = sample_df.to_string(index=False)

    # Full data
    full_text = df.to_string(index=False)

    text = "\n".join(schema_lines) + "\n\n--- Sample (first 5 rows) ---\n" + sample_text + "\n\n--- Full Data ---\n" + full_text

    result: Dict[str, Any] = {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": ext.lstrip("."),
        "page_count": None,
        "row_count": row_count,
        "columns": json.dumps(columns[:50]),
        "schema": json.dumps(type_map),
    }
    if truncated:
        result["truncated"] = True

    return result


@register_parser([".html", ".htm"])
def parse_html(file_path: str) -> Dict[str, Any]:
    """Parse HTML files, stripping tags to extract readable text."""
    path = Path(file_path)
    raw = path.read_text(encoding="utf-8", errors="replace")

    try:
        from html.parser import HTMLParser

        class _TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self._parts: List[str] = []
                self._skip = False

            def handle_starttag(self, tag, attrs):
                self._skip = tag.lower() in ("script", "style", "noscript")

            def handle_endtag(self, tag):
                if tag.lower() in ("script", "style", "noscript"):
                    self._skip = False

            def handle_data(self, data):
                if not self._skip:
                    stripped = data.strip()
                    if stripped:
                        self._parts.append(stripped)

        extractor = _TextExtractor()
        extractor.feed(raw)
        text = "\n".join(extractor._parts)
    except Exception:
        # Fallback: return raw if parsing fails
        text = raw

    return {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": "html",
        "page_count": None,
    }


@register_parser([
    ".txt", ".md", ".rst", ".log",
    ".py", ".js", ".ts", ".jsx", ".tsx",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg",
    ".sh", ".bash",
    ".xml",
    ".java", ".go", ".rs", ".rb", ".cpp", ".c", ".h", ".cs",
    ".sql", ".r", ".swift", ".kt",
])
def parse_text(file_path: str) -> Dict[str, Any]:
    path = Path(file_path)

    # Binary file detection: check first 512 bytes for null bytes
    try:
        with open(file_path, "rb") as f:
            sample = f.read(512)
        if b"\x00" in sample:
            raise ValueError(
                f"File '{path.name}' appears to be a binary file "
                f"(null bytes detected). Only text files are supported."
            )
    except ValueError:
        raise
    except Exception:
        pass  # proceed with text read if binary check fails

    text = path.read_text(encoding="utf-8", errors="replace")
    return {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": path.suffix.lstrip("."),
        "page_count": None,
    }


# ---------------------------------------------------------------------------
# Phase 8B: New community parsers
# ---------------------------------------------------------------------------

@register_parser([".eml"])
def parse_eml(file_path: str) -> Dict[str, Any]:
    """
    Parse .eml email files — extract headers, body text, and attachment list.

    Handles MIME multipart messages, preferring text/plain over text/html.
    Attachment filenames are listed as metadata (not extracted).
    """
    import email
    import email.policy
    from email import message_from_bytes

    path = Path(file_path)
    raw = path.read_bytes()

    try:
        msg = message_from_bytes(raw, policy=email.policy.default)
    except Exception as e:
        raise ValueError(
            f"Failed to parse email '{path.name}': {e}. "
            f"File may not be a valid .eml file."
        ) from e

    # Extract headers
    headers = {}
    for key in ("From", "To", "Cc", "Subject", "Date", "Message-ID"):
        val = msg.get(key, "")
        if val:
            headers[key] = str(val)

    header_text = "\n".join(f"{k}: {v}" for k, v in headers.items())

    # Extract body — prefer text/plain, fall back to text/html (stripped)
    body = ""
    attachments = []

    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            disposition = str(part.get("Content-Disposition", ""))

            # Track attachments
            if "attachment" in disposition:
                att_name = part.get_filename() or "(unnamed)"
                att_size = len(part.get_payload(decode=True) or b"")
                attachments.append(f"{att_name} ({att_size} bytes)")
                continue

            if content_type == "text/plain" and not body:
                payload = part.get_payload(decode=True)
                if payload:
                    body = payload.decode("utf-8", errors="replace")
            elif content_type == "text/html" and not body:
                payload = part.get_payload(decode=True)
                if payload:
                    html = payload.decode("utf-8", errors="replace")
                    body = _strip_html_tags(html)
    else:
        content_type = msg.get_content_type()
        payload = msg.get_payload(decode=True)
        if payload:
            raw_text = payload.decode("utf-8", errors="replace")
            if content_type == "text/html":
                body = _strip_html_tags(raw_text)
            else:
                body = raw_text

    # Compose full text
    parts = [header_text]
    if body:
        parts.append(f"\n--- Body ---\n{body.strip()}")
    if attachments:
        parts.append(f"\n--- Attachments ({len(attachments)}) ---\n" + "\n".join(attachments))

    text = "\n".join(parts)

    return {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": "eml",
        "page_count": None,
        "attachment_count": len(attachments),
        "subject": headers.get("Subject", ""),
    }


@register_parser([".mbox"])
def parse_mbox(file_path: str) -> Dict[str, Any]:
    """
    Parse .mbox mailbox files — extract all messages as individual sections.

    Limits to first 100 messages to prevent memory issues with large archives.
    """
    import mailbox

    path = Path(file_path)
    try:
        mbox = mailbox.mbox(file_path)
    except Exception as e:
        raise ValueError(
            f"Failed to parse mbox '{path.name}': {e}. "
            f"File may not be a valid .mbox file."
        ) from e

    messages = []
    max_messages = 100
    total_count = 0

    for msg in mbox:
        total_count += 1
        if len(messages) >= max_messages:
            continue  # count but don't extract

        subject = msg.get("Subject", "(no subject)")
        from_addr = msg.get("From", "")
        date = msg.get("Date", "")

        # Get body text
        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        body = payload.decode("utf-8", errors="replace")
                        break
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                body = payload.decode("utf-8", errors="replace")

        header = f"From: {from_addr}\nDate: {date}\nSubject: {subject}"
        messages.append(f"{header}\n\n{body.strip()}")

    mbox.close()

    sep = "\n\n" + "=" * 60 + "\n\n"
    text = sep.join(messages)
    if total_count > max_messages:
        text += f"\n\n[... {total_count - max_messages} more messages truncated ...]"

    return {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": "mbox",
        "page_count": total_count,
    }


@register_parser([".epub"])
def parse_epub(file_path: str) -> Dict[str, Any]:
    """
    Parse EPUB e-books — extract text content from XHTML chapters.

    Uses zipfile to read the EPUB container (which is a ZIP archive),
    locates content documents via the OPF manifest, and extracts text
    from each XHTML chapter in reading order.

    No external dependencies required — uses stdlib zipfile + html.parser.
    """
    import xml.etree.ElementTree as ET
    import zipfile

    path = Path(file_path)

    try:
        zf = zipfile.ZipFile(file_path, "r")
    except Exception as e:
        raise ValueError(
            f"Failed to read EPUB '{path.name}': {e}. "
            f"File may be corrupted or not a valid .epub file."
        ) from e

    chapters = []
    title = ""

    try:
        # Find the OPF file via container.xml
        try:
            container_xml = zf.read("META-INF/container.xml")
            container_root = ET.fromstring(container_xml)
            # Handle namespace
            ns = {"c": "urn:oasis:names:tc:opendocument:xmlns:container"}
            rootfile = container_root.find(".//c:rootfile", ns)
            opf_path = rootfile.get("full-path", "") if rootfile is not None else ""
        except Exception:
            # Fallback: look for .opf file
            opf_path = ""
            for name in zf.namelist():
                if name.endswith(".opf"):
                    opf_path = name
                    break

        if not opf_path:
            raise ValueError(f"EPUB '{path.name}': cannot find OPF manifest")

        opf_dir = str(Path(opf_path).parent)
        if opf_dir == ".":
            opf_dir = ""

        # Parse OPF to get spine (reading order) and manifest
        opf_data = zf.read(opf_path)
        opf_root = ET.fromstring(opf_data)

        # Extract namespaces dynamically
        opf_ns = {"opf": "http://www.idpf.org/2007/opf", "dc": "http://purl.org/dc/elements/1.1/"}

        # Get title
        title_el = opf_root.find(".//dc:title", opf_ns)
        if title_el is not None and title_el.text:
            title = title_el.text.strip()

        # Build manifest id → href mapping
        manifest = {}
        for item in opf_root.findall(".//opf:manifest/opf:item", opf_ns):
            item_id = item.get("id", "")
            href = item.get("href", "")
            media_type = item.get("media-type", "")
            if item_id and href:
                manifest[item_id] = {"href": href, "media_type": media_type}

        # Get spine (reading order)
        spine_refs = []
        for itemref in opf_root.findall(".//opf:spine/opf:itemref", opf_ns):
            idref = itemref.get("idref", "")
            if idref and idref in manifest:
                spine_refs.append(manifest[idref])

        # If no spine found, fall back to all XHTML items from manifest
        if not spine_refs:
            spine_refs = [
                info for info in manifest.values()
                if info["media_type"] in ("application/xhtml+xml", "text/html")
            ]

        # Extract text from each chapter
        for ref in spine_refs:
            href = ref["href"]
            # Resolve relative path
            if opf_dir:
                full_path = f"{opf_dir}/{href}"
            else:
                full_path = href

            try:
                content = zf.read(full_path).decode("utf-8", errors="replace")
                chapter_text = _strip_html_tags(content)
                if chapter_text.strip():
                    chapters.append(chapter_text.strip())
            except (KeyError, Exception) as e:
                logger.debug(f"EPUB: skipping {href}: {e}")

    finally:
        zf.close()

    if not chapters:
        raise ValueError(
            f"No text content found in EPUB '{path.name}'. "
            f"File may contain only images or be DRM-protected."
        )

    header = f"Title: {title}\n\n" if title else ""
    text = header + "\n\n---\n\n".join(chapters)

    return {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": "epub",
        "page_count": len(chapters),
        "title": title,
    }


@register_parser([".rtf"])
def parse_rtf(file_path: str) -> Dict[str, Any]:
    """
    Parse RTF (Rich Text Format) files — extract plain text content.

    Uses a lightweight state-machine RTF stripper (no external deps).
    Handles control words, groups, Unicode escapes, and hex characters.
    """
    path = Path(file_path)
    raw = path.read_bytes()

    try:
        text = _strip_rtf(raw)
    except Exception as e:
        raise ValueError(
            f"Failed to parse RTF '{path.name}': {e}. "
            f"File may be corrupted or not a valid .rtf file."
        ) from e

    if not text.strip():
        raise ValueError(f"No text content found in RTF '{path.name}'.")

    return {
        "text": text[:_MAX_TEXT_CHARS],
        "file_type": "rtf",
        "page_count": None,
    }


# ---------------------------------------------------------------------------
# Shared utility functions
# ---------------------------------------------------------------------------

def _strip_html_tags(html: str) -> str:
    """Strip HTML tags and return plain text. Lightweight, no external deps."""
    from html.parser import HTMLParser

    class _Stripper(HTMLParser):
        def __init__(self):
            super().__init__()
            self._parts: List[str] = []
            self._skip = False

        def handle_starttag(self, tag, attrs):
            self._skip = tag.lower() in ("script", "style", "noscript")
            # Add newline for block-level elements
            if tag.lower() in ("p", "div", "br", "h1", "h2", "h3", "h4", "h5", "h6", "li", "tr"):
                self._parts.append("\n")

        def handle_endtag(self, tag):
            if tag.lower() in ("script", "style", "noscript"):
                self._skip = False

        def handle_data(self, data):
            if not self._skip:
                self._parts.append(data)

    try:
        stripper = _Stripper()
        stripper.feed(html)
        return "".join(stripper._parts).strip()
    except Exception:
        # Last resort: regex strip
        import re
        return re.sub(r"<[^>]+>", " ", html).strip()


def _strip_rtf(raw: bytes) -> str:
    """
    Strip RTF control codes and return plain text.

    Lightweight state-machine approach — handles:
    - Control words (\\word) and control symbols (\\*)
    - Group nesting ({}), skipping destination groups
    - Unicode escapes (\\uN) and hex chars (\\'XX)
    - Special characters: \\par → newline, \\tab → tab
    """
    import re

    text = raw.decode("ascii", errors="replace")

    # Destinations to skip (contain no user-visible text)
    _skip_destinations = {
        "fonttbl", "colortbl", "stylesheet", "info", "pict",
        "header", "footer", "headerl", "headerr", "footerl", "footerr",
        "headerf", "footerf", "object", "objdata", "datafield",
        "fldinst", "themedata", "colorschememapping", "datastore",
        "latentstyles", "generator",
    }

    output = []
    i = 0
    length = len(text)
    group_depth = 0
    skip_depth = 0  # depth at which we started skipping

    while i < length:
        ch = text[i]

        if ch == "{":
            group_depth += 1
            # Check if this group starts with a skippable destination
            if i + 1 < length and text[i + 1] == "\\":
                # Look ahead for control word
                m = re.match(r"\\(\*\\)?([a-z]+)", text[i + 1:i + 40])
                if m:
                    word = m.group(2)
                    if word in _skip_destinations:
                        skip_depth = group_depth
            i += 1
            continue

        if ch == "}":
            if group_depth == skip_depth:
                skip_depth = 0
            group_depth -= 1
            i += 1
            continue

        if skip_depth > 0:
            i += 1
            continue

        if ch == "\\":
            i += 1
            if i >= length:
                break

            next_ch = text[i]

            # Hex char: \'XX
            if next_ch == "'":
                if i + 2 < length:
                    hex_val = text[i + 1:i + 3]
                    try:
                        output.append(chr(int(hex_val, 16)))
                    except ValueError:
                        pass
                    i += 3
                    continue
                i += 1
                continue

            # Unicode: \uN followed by replacement char
            if next_ch == "u":
                m = re.match(r"(-?\d+)", text[i + 1:i + 8])
                if m:
                    code = int(m.group(1))
                    if code < 0:
                        code += 65536
                    try:
                        output.append(chr(code))
                    except ValueError:
                        pass
                    i += 1 + len(m.group(1))
                    # Skip replacement character
                    if i < length and text[i] == " ":
                        i += 1
                    continue
                i += 1
                continue

            # Control word
            if next_ch.isalpha():
                m = re.match(r"([a-z]+)(-?\d+)?", text[i:i + 30])
                if m:
                    word = m.group(1)
                    i += len(m.group(0))
                    # Skip optional trailing space
                    if i < length and text[i] == " ":
                        i += 1
                    # Map control words to text
                    if word == "par" or word == "line":
                        output.append("\n")
                    elif word == "tab":
                        output.append("\t")
                    elif word in ("lquote", "rquote"):
                        output.append("'")
                    elif word in ("ldblquote", "rdblquote"):
                        output.append('"')
                    elif word in ("emdash", "endash"):
                        output.append("—" if word == "emdash" else "–")
                    elif word == "bullet":
                        output.append("•")
                    continue
                i += 1
                continue

            # Control symbol (\\ \{ \} etc.)
            if next_ch in ("\\", "{", "}"):
                output.append(next_ch)
                i += 1
                continue

            # Other control symbol — skip
            i += 1
            continue

        # Regular character
        if ch in ("\r", "\n"):
            i += 1
            continue
        output.append(ch)
        i += 1

    result = "".join(output)
    # Clean up excessive whitespace
    result = re.sub(r"\n{3,}", "\n\n", result)
    result = re.sub(r"[ \t]+", " ", result)
    return result.strip()
